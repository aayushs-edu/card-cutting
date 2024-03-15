import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
import os

from parsing import Card, Body

year = '2013OpenEv'

datapath = os.path.join(os.getcwd(), 'selenium_downloads')
dirpath = os.path.join(os.getcwd(), f'selenium_downloads\\{year}')


def parse(doc : docx.Document):
    paras = doc.paragraphs
    num_paras = len(doc.paragraphs)
    print('Paragraphs to parse: ', num_paras)
    headings = []
    title = paras[0].text
    filepath = os.path.join(dirpath, title)
    os.mkdir(filepath)
    for i in range(1, num_paras):
        # Create new document for each card
        text = paras[i].text
        style = paras[i].style.name
        if style.startswith('Heading 2'):
            h = text[1:].replace('/', '-')
            os.mkdir(os.path.join(filepath, h))
            headings.append(h)
        elif style.startswith('Heading 3'):
            print('New card: ', text)
            heading = paras[i].text.replace('/', '-')
            new_doc = Document()
            i+=1
            while not paras[i].style.name.startswith('Heading 3'):
                print('Adding paragraph: ', paras[i].text, 'Para style: ', paras[i].style.name)
                new_p = new_doc.add_paragraph()
                for run in paras[i].runs:
                    new_run = new_p.add_run(text=run.text)
                    if run.style.name == 'Style Bold Underline':
                        new_run.underline = True
                    elif run.style.name == 'Emphasis':
                        new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    elif run.font.size == 101600:
                        new_run.font.size = Pt(6)
                    elif run.font.size == 76200:
                        new_run.font.size = Pt(6)
                    elif run.bold:
                        new_run.bold = True       
                i+=1
                if i >= num_paras: break
            print(f'Saving {heading}.docx to ', os.path.join(filepath, heading+'.docx'))
            new_doc.save(os.path.join(filepath, headings[-1], heading+'.docx'))
            i-=1
            yield new_doc

def get_cards(doc : docx.Document):
    paras = doc.paragraphs
    for para in paras:
        if para.style.name.startswith('Heading'):
            yield para.text

def doc_struct(doc : docx.Document):
    result : dict = {}
    paras = doc.paragraphs
    num_paras = len(paras)

    for i in range(num_paras):
        style = paras[i].style.name
        if style == 'Heading 1': continue
        if style == 'Heading 2':
            print('NEW SECTION...')
            heading = paras[i].text
            cards : dict = {}
            i += 1
            while True:
                if i >= num_paras - 1: break
                if paras[i].style.name == 'Heading 2': 
                    print('SECTION ENDED')
                    break
                if paras[i].style.name == 'Heading 3':
                    if not paras[i].text: 
                        i += 1
                        continue
                    print('New card...')
                    card_title = paras[i].text
                    contents : list[docx.text.paragraph.Paragraph] = []
                    i += 1
                    while True:
                        print('Appending paragraph: ', paras[i].text)
                        contents.append(paras[i])
                        if i+1 >= num_paras: break
                        if paras[i+1].style.name.startswith('Heading'):
                            print('Card ended')
                            break
                        i += 1
                    cards[card_title] = contents
                i += 1
            result[heading] = cards
    return result

def split_docs(doc_struct : dict, doc_title):

    filepath = os.path.join(dirpath, doc_title)
    os.mkdir(filepath)

    for section, cards in doc_struct.items():
        h = section[1:].replace('/', '-')
        os.mkdir(os.path.join(filepath, h))
        for card_name, content in cards.items():
            d = Document()
            for para in content:
                p = d.add_paragraph()
                for run in para.runs:
                    new_run = p.add_run(text=run.text)
                    if run.style.name == 'Style Bold Underline':
                        new_run.underline = True
                    elif run.style.name == 'Emphasis':
                        new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    elif run.font.size == 101600:
                        new_run.font.size = Pt(6)
                    elif run.font.size == 76200:
                        new_run.font.size = Pt(6)
                    elif run.bold:
                        new_run.bold = True
            d.save(os.path.join(filepath, h, card_name.replace('/', '-').replace(':', '-')+'.docx'))
            print(f'Saving {card_name}.docx')


def parse2(doc : docx.Document, filename : str, savePath):

    paras = doc.paragraphs
    os.mkdir(os.path.join(savePath, filename))

    num_paras = len(paras)
    for i in range(num_paras):
        if paras[i].style.name == 'Heading 3' and paras[i].text:
            new_doc = Document()
            title = paras[i].text
            i += 1
            while True:
                # print('Appending paragraph: ', paras[i].text)
                p = new_doc.add_paragraph()
                for run in paras[i].runs:
                    new_run = p.add_run(run.text)
                    if not run.style: continue
                    if run.style.name == 'Style Bold Underline':
                        new_run.underline = True
                    if run.style.name == 'Emphasis' or run.font.highlight_color != None:
                        new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    if run.font.size == 101600:
                        new_run.font.size = Pt(6)
                    if run.font.size == 76200:
                        new_run.font.size = Pt(6)
                    if run.bold or 'Bold' in run.style.name or paras[i].style.name == 'Heading 4':
                        new_run.bold = True
                if i+1 >= num_paras: break
                if paras[i+1].style.name == 'Heading 2' or paras[i+1].style.name == 'Heading 3':
                    # print('Card ended')
                    break
                i += 1
            new_doc.save(os.path.join(savePath, filename, title.replace("\t", '').replace('"', '').replace("\n", '').replace('/', '-').replace(':', '-').replace('?','').replace('*', '').replace('>', 'greater than')+'.docx'))
            print(f'Saving {title}.docx')
            

# doc_struct = doc_struct(document)
# print(doc_struct.keys())

# print(list(para.text for para in doc_struct['*Latin American Relations']['1nc--Poverty']))

def parseYear(year : str):
    savePath = os.path.join(datapath, year[:4])
    getpath = os.path.join(datapath, year)
    if not os.path.exists(savePath): os.mkdir(savePath)
    for file in os.listdir(getpath):
        if os.path.exists(os.path.join(savePath, file)) or not file.endswith('.docx'): continue
        try: document : docx.Document = Document(os.path.join(getpath, file))
        except Exception: continue
        parse2(document, file, savePath)

# parseYear('2013OpenEv')

# print(list(get_cards(document)))
# print(list(parse(document)))

yearlist = os.path.join(datapath, '2013')    
cardlist = os.path.join(yearlist, os.listdir(yearlist)[0]) 
currCard = os.path.join(cardlist, os.listdir(cardlist)[0])

cardFile = Document(currCard)
card1 = Card(os.path.basename(currCard))
print(card1.title)
paras = cardFile.paragraphs
card1.set_tagline(paras[0])
card1.add_evidence()
