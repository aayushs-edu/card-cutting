from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
import os
import math
import polars as pl
import string
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
import numpy as np
import re
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

from synonyms import data

SAVE_DIR = os.getcwd()

regex = re.compile('[^a-zA-Z]')

# Selenium setup (need selenium to get article paragraph by paragraph)

chrome_options = Options()
chrome_options.add_argument("--window-size=1920,1080")
chrome_options.add_argument("--disable-dev-shm-usage")
chrome_options.headless = True
service = Service(os.getenv('driverpath'))
driver = webdriver.Chrome(service=service, options=chrome_options)

# Stopwords used to filter taglines

stopwords = ["copyright", "0o", "0s", "3a", "3b", "3d", "6b", "6o", "a", "a1", "a2", "a3", "a4", "ab", 
             "able", "about", "above", "abst", "ac", "accordance", "according", "accordingly", 
             "across", "act", "actually", "ad", "added", "adj", "ae", "af", "affected", 
             "affecting", "affects", "after", "afterwards", "ag", "again", "against", "ah", 
             "ain", "ain't", "aj", "al", "all", "allow", "allows", "almost", "alone", "along", 
             "already", "also", "although", "always", "am", "among", "amongst", "amoungst", 
             "amount", "an", "and", "announce", "another", "any", "anybody", "anyhow", "anymore", 
             "anyone", "anything", "anyway", "anyways", "anywhere", "ao", "ap", "apart", 
             "apparently", "appear", "appreciate", "appropriate", "approximately", "ar", "are", 
             "aren", "arent", "aren't", "arise", "around", "as", "a's", "aside", "ask", "asking", 
             "associated", "at", "au", "auth", "av", "available", "aw", "away", "awfully", "ax", 
             "ay", "az", "b", "b1", "b2", "b3", "ba", "back", "bc", "bd", "be", "became", "because", "become", "becomes", "becoming", "been", "before", "beforehand", "begin", "beginning", "beginnings", "begins", "behind", "being", "believe", "below", "beside", "besides", "best", "better", "between", "beyond", "bi", "bill", "biol", "bj", "bk", "bl", "bn", "both", "bottom", "bp", "br", "brief", "briefly", "bs", "bt", "bu", "but", "bx", "by", "c", "c1", "c2", "c3", "ca", "call", "came", "can", "cannot", "cant", "can't", "cause", "causes", "cc", "cd", "ce", "certain", "certainly", "cf", "cg", "ch", "changes", "ci", "cit", "cj", "cl", "clearly", "cm", "c'mon", "cn", "co", "com", "come", "comes", "con", "concerning", "consequently", "consider", "considering", "contain", "containing", "contains", "corresponding", "could", "couldn", "couldnt", "couldn't", "course", "cp", "cq", "cr", "cry", "cs", "c's", "ct", "cu", "currently", "cv", "cx", "cy", "cz", "d", "d2", "da", "date", "dc", "dd", "de", "definitely", "describe", "described", "despite", "detail", "df", "di", "did", "didn", "didn't", "different", "dj", "dk", "dl", "do", "does", "doesn", "doesn't", "doing", "don", "done", "don't", "down", "downwards", "dp", "dr", "ds", "dt", "du", "due", "during", "dx", "dy", "e", "e2", "e3", "ea", "each", "ec", "ed", "edu", "ee", "ef", "effect", "eg", "ei", "eight", "eighty", "either", "ej", "el", "eleven", "else", "elsewhere", "em", "empty", "en", "end", "ending", "enough", "entirely", "eo", "ep", "eq", "er", "es", "especially", "est", "et", "et-al", "etc", "eu", "ev", "even", "ever", "every", "everybody", "everyone", "everything", "everywhere", "ex", "exactly", "example", "except", "ey", "f", "f2", "fa", "far", "fc", "few", "ff", "fi", "fifteen", "fifth", "fify", "fill", "find", "fire", "first", "five", "fix", "fj", "fl", "fn", "fo", "followed", "following", "follows", "for", "former", "formerly", "forth", "forty", "found", "four", "fr", "from", "front", "fs", "ft", "fu", "full", "further", "furthermore", "fy", "g", "ga", "gave", "ge", "get", "gets", "getting", "gi", "give", "given", "gives", "giving", "gj", "gl", "go", "goes", "going", "gone", "got", "gotten", "gr", "greetings", "gs", "gy", "h", "h2", "h3", "had", "hadn", "hadn't", "happens", "hardly", "has", "hasn", "hasnt", "hasn't", "have", "haven", "haven't", "having", "he", "hed", "he'd", "he'll", "hello", "help", "hence", "her", "here", "hereafter", "hereby", "herein", "heres", "here's", "hereupon", "hers", "herself", "hes", "he's", "hh", "hi", "hid", "him", "himself", "his", "hither", "hj", "ho", "home", "hopefully", "how", "howbeit", "however", "how's", "hr", "hs", "http", "hu", "hundred", "hy", "i", "i2", "i3", "i4", "i6", "i7", "i8", "ia", "ib", "ibid", "ic", "id", "i'd", "ie", "if", "ig", "ignored", "ih", "ii", "ij", "il", "i'll", "im", "i'm", "immediate", "immediately", "importance", "important", "in", "inasmuch", "inc", "indeed", "index", "indicate", "indicated", "indicates", "information", "inner", "insofar", "instead", "interest", "into", "invention", "inward", "io", "ip", "iq", "ir", "is", "isn", "isn't", "it", "itd", "it'd", "it'll", "its", "it's", "itself", "iv", "i've", "ix", "iy", "iz", "j", "jj", "jr", "js", "jt", "ju", "just", "k", "ke", "keep", "keeps", "kept", "kg", "kj", "km", "know", "known", "knows", "ko", "l", "l2", "la", "largely", "last", "lately", "later", "latter", "latterly", "lb", "lc", "le", "least", "les", "less", "lest", "let", "lets", "let's", "lf", "like", "liked", "likely", "line", "little", "lj", "ll", "ll", "ln", "lo", "look", "looking", "looks", "los", "lr", "ls", "lt", "ltd", "m", "m2", "ma", "made", "mainly", "make", "makes", "many", "may", "maybe", "me", "mean", "means", "meantime", "meanwhile", "merely", "mg", "might", "mightn", "mightn't", "mill", "million", "mine", "miss", "ml", "mn", "mo", "more", "moreover", "most", "mostly", "move", "mr", "mrs", "ms", "mt", "mu", "much", "mug", "must", "mustn", "mustn't", "my", "myself", "n", "n2", "na", "name", "namely", "nay", "nc", "nd", "ne", "near", "nearly", "necessarily", "necessary", "need", "needn", "needn't", "needs", "neither", "never", "nevertheless", "new", "next", "ng", "ni", "nine", "ninety", "nj", "nl", "nn", "no", "nobody", "non", "none", "nonetheless", "noone", "nor", "normally", "nos", "not", "noted", "nothing", "novel", "now", "nowhere", "nr", "ns", "nt", "ny", "o", "oa", "ob", "obtain", "obtained", "obviously", "oc", "od", "of", "off", "often", "og", "oh", "oi", "oj", "ok", "okay", "ol", "old", "om", "omitted", "on", "once", "one", "ones", "only", "onto", "oo", "op", "oq", "or", "ord", "os", "ot", "other", "others", "otherwise", "ou", "ought", "our", "ours", "ourselves", "out", "outside", "over", "overall", "ow", "owing", "own", "ox", "oz", "p", "p1", "p2", "p3", "page", "pagecount", "pages", "par", "part", "particular", "particularly", "pas", "past", "pc", "pd", "pe", "per", "perhaps", "pf", "ph", "pi", "pj", "pk", "pl", "placed", "please", "plus", "pm", "pn", "po", "poorly", "possible", "possibly", "potentially", "pp", "pq", "pr", "predominantly", "present", "presumably", "previously", "primarily", "probably", "promptly", "proud", "provides", "ps", "pt", "pu", "put", "py", "q", "qj", "qu", "que", "quickly", "quite", "qv", "r", "r2", "ra", "ran", "rather", "rc", "rd", "re", "readily", "really", "reasonably", "recent", "recently", "ref", "refs", "regarding", "regardless", "regards", "related", "relatively", "research", "research-articl", "respectively", "resulted", "resulting", "results", "rf", "rh", "ri", "right", "rj", "rl", "rm", "rn", "ro", "rq", "rr", "rs", "rt", "ru", "run", "rv", "ry", "s", "s2", "sa", "said", "same", "saw", "say", "saying", "says", "sc", "sd", "se", "sec", "second", "secondly", "section", "see", "seeing", "seem", "seemed", "seeming", "seems", "seen", "self", "selves", "sensible", "sent", "serious", "seriously", "seven", "several", "sf", "shall", "shan", "shan't", "she", "shed", "she'd", "she'll", "shes", "she's", "should", "shouldn", "shouldn't", "should've", "show", "showed", "shown", "showns", "shows", "si", "side", "significant", "significantly", "similar", "similarly", "since", "sincere", "six", "sixty", "sj", "sl", "slightly", "sm", "sn", "so", "some", "somebody", "somehow", "someone", "somethan", "something", "sometime", "sometimes", "somewhat", "somewhere", "soon", "sorry", "sp", "specifically", "specified", "specify", "specifying", "sq", "sr", "ss", "st", "still", "stop", "strongly", "sub", "substantially", "successfully", "such", "sufficiently", "suggest", "sup", "sure", "sy", "system", "sz", "t", "t1", "t2", "t3", "take", "taken", "taking", "tb", "tc", "td", "te", "tell", "ten", "tends", "tf", "th", "than", "thank", "thanks", "thanx", "that", "that'll", "thats", "that's", "that've", "the", "their", "theirs", "them", "themselves", "then", "thence", "there", "thereafter", "thereby", "thered", "therefore", "therein", "there'll", "thereof", "therere", "theres", "there's", "thereto", "thereupon", "there've", "these", "they", "theyd", "they'd", "they'll", "theyre", "they're", "they've", "thickv", "thin", "think", "third", "this", "thorough", "thoroughly", "those", "thou", "though", "thoughh", "thousand", "three", "throug", "through", "throughout", "thru", "thus", "ti", "til", "tip", "tj", "tl", "tm", "tn", "to", "together", "too", "took", "top", "toward", "towards", "tp", "tq", "tr", "tried", "tries", "truly", "try", "trying", "ts", "t's", "tt", "tv", "twelve", "twenty", "twice", "two", "tx", "u", "u201d", "ue", "ui", "uj", "uk", "um", "un", "under", "unfortunately", "unless", "unlike", "unlikely", "until", "unto", "uo", "up", "upon", "ups", "ur", "us", "use", "used", "useful", "usefully", "usefulness", "uses", "using", "usually", "ut", "v", "va", "value", "various", "vd", "ve", "ve", "very", "via", "viz", "vj", "vo", "vol", "vols", "volumtype", "vq", "vs", "vt", "vu", "w", "wa", "want", "wants", "was", "wasn", "wasnt", "wasn't", "way", "we", "wed", "we'd", "welcome", "well", "we'll", "well-b", "went", "were", "we're", "weren", "werent", "weren't", "we've", "what", "whatever", "what'll", "whats", "what's", "when", "whence", "whenever", "when's", "where", "whereafter", "whereas", "whereby", "wherein", "wheres", "where's", "whereupon", "wherever", "whether", "which", "while", "whim", "whither", "who", "whod", "whoever", "whole", "who'll", "whom", "whomever", "whos", "who's", "whose", "why", "why's", "wi", "widely", "will", "willing", "wish", "with", "within", "without", "wo", "won", "wonder", "wont", "won't", "words", "world", "would", "wouldn", "wouldnt", "wouldn't", "www", "x", "x1", "x2", "x3", "xf", "xi", "xj", "xk", "xl", "xn", "xo", "xs", "xt", "xv", "xx", "y", "y2", "yes", "yet", "yj", "yl", "you", "youd", "you'd", "you'll", "your", "youre", "you're", "yours", "yourself", "yourselves", "you've", "yr", "ys", "yt", "z", "zero", "zi", "zz"]

################################################ FUNCTIONS NEEDED ###############################################################

def tfidf(term : str, paragraphs : list[str]):
    nParas = len(paragraphs)

    # tf
    tfs = [math.log1p(p.lower().count(term)/len(p.split(' '))) * 10 for p in paragraphs]
    # print(f'Tfs for {term} in article: ', tfs)

    # idf
    parasWithTerm = 0
    for p in paragraphs:
        parasWithTerm += 1 * (term in p)
    # print(f'Paras with term: {parasWithTerm}')
    if parasWithTerm == 0: 
        return 0, 0, 0, 0

    idf = math.log2(nParas / parasWithTerm)
    # print(f'Idf for {term} in article: ', idf)

    tfidfs = [tf * idf for tf in tfs]
    # print(f'TFIDFS for {term} in article: ', tfidfs)

    relevantParas = { para : val for (para, val) in zip(paragraphs, tfidfs) if val != 0}
    
    return tfs, idf, tfidfs, relevantParas

def weightedTfidf(term : tuple[str, float], paragraphs : list[str]):
    nParas = len(paragraphs)

    # tf
    tfs = [math.log1p(p.lower().count(term[0])/len(p.split(' '))) * 10 for p in paragraphs]
    # print(f'Tfs for {term} in article: ', tfs)

    # idf
    parasWithTerm = 0
    for p in paragraphs:
        parasWithTerm += 1 * (term[0] in p)
    # print(f'Paras with term: {parasWithTerm}')
    if parasWithTerm == 0: 
        return 0, 0, 0, 0

    idf = math.log2(nParas / parasWithTerm)
    # print(f'Idf for {term} in article: ', idf)

    tfidfs = [tf * idf for tf in tfs]
    # print(f'TFIDFS for {term} in article: ', tfidfs)

    relevantParas = { para : val * term[1] for (para, val) in zip(paragraphs, tfidfs) if val != 0}
    
    return tfs, idf, tfidfs, relevantParas

# Looping through tagline words and using tfidf

def labeledRanking(tagline_words : list[str], paragraphs : list[str]):
    finalParas = {}

    for word in tagline_words:
        _, _, _, relevantParas = tfidf(word, paragraphs)
        if relevantParas:
            finalParas[word] = relevantParas
    return finalParas

# For unlabeled ranking, we need extra logic to maintain order of sentences as they were in the article
def unlabeledRanking(tagline_words : str | list[str] | dict[ str, float ], paragraphs : list[str]) -> dict[str, float]:
    rankedParas = {}
    if isinstance(tagline_words, str):
        tagline_words = tagline_words.split(' ')

    for word in tagline_words:
        _, _, _, relevantParas = tfidf(word, paragraphs)
        if relevantParas:
            rankedParas.update(relevantParas) 
    
    return sortSentences(rankedParas, paragraphs)

# For unlabeled ranking, we need extra logic to maintain order of sentences as they were in the article
def weightedRanking(tagline_words : dict[ str, float ], paragraphs : list[str]) -> dict[str, float]:
    rankedParas = {}

    for word, val in tagline_words.items():
        _, _, _, relevantParas = weightedTfidf((word, val), paragraphs)
        if relevantParas:
            rankedParas.update(relevantParas) 
    
    return sortSentences(rankedParas, paragraphs)

def sortSentences(sentences : dict[str : float], articleParas : list[str]) -> dict[str : float]:
    d = {}
    for pair in sentences.items():
        loc = articleParas.index(pair[0])
        d[loc] = pair

    final = {}
    for pair in dict(sorted(d.items())).values():
        final[pair[0]] = pair[1]
    return final
    
def makeDoc(tagline : str, site : str, articleTitle : str, date : str, paraDict : dict[ str : float ], docName : str):
    doc = Document()
    ranks = list(paraDict.values())

    # Add tagline
    tl = doc.add_paragraph()
    tl.add_run(tagline).bold = True

    # Add source using link
    source = f' "{articleTitle}", {date}'
    src = doc.add_paragraph()
    site_name = src.add_run(site)
    site_name.bold = True
    src.add_run(source)

    body = doc.add_paragraph()
    for para, val in paraDict.items():
        # print(para)
        if val > 0:
            run = body.add_run(para)
            if val >= np.percentile(ranks, 60):
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                run.bold = True
                run.underline = True
            elif val >= np.percentile(ranks, 40):
                run.bold = True
                run.underline = True
            elif val >= np.percentile(ranks, 30):
                run.underline = True
            else:
                run.font.size = Pt(6)
            
    doc.save(os.path.join(SAVE_DIR, f'{docName}.docx'))

def removePunc(s):
    return s.translate(str.maketrans('', '', string.punctuation))

def parseTagline(tagline : str) -> list[str]:
    tagline = removePunc(tagline)
    words = []
    for word in tagline.split(' '):
        if word.lower() not in stopwords and not word.isnumeric() and word.isalpha() and len(word) > 3:
            words.append(word.lower())
    return words

def jaccard(str1 : str, str2 : str):
    str1words = set(removePunc(str1).split(' '))
    str2words = set(removePunc(str2).split(' '))
    union = str1words.union(str2words)
    intersection = str1words.intersection(str2words)
    return len(intersection) / len(union)

def findSimilarTaglinesTFIDF(tagline : str, dataset : list[str]):
    tagline_words = parseTagline(tagline)
    all_synonym_tls : dict[str, float] = {}
    for word in tagline_words:
        _, _, _, similarTaglines = tfidf(word.lower(), dataset)
        if similarTaglines:
            all_synonym_tls.update(similarTaglines)
    
    return all_synonym_tls

def findSimilarTaglinesJaccard(tagline : str, dataset : list[str]):
    res : dict[ str, float ] = {}
    for tl in dataset:
        res[tl] = jaccard(tagline, tl)
    return res

def abovePercentile(d : dict[str, float], percentile : int):
    perc = np.percentile(list(d.values()), percentile)
    # print(f'Percentile for {list(d.values())}: {perc}')
    new_d = { sentence : val for sentence, val in d.items() if val >= perc }
    return new_d

def generateTxt(vals : dict[ str, float ], fileName : str):
    f = open(fileName+'.txt', 'x')
    for sentence, val in vals.items():
        f.write(sentence + '\t' + str(val) + '\n')

def uniqueTaglineWords(similarTaglines : list[str]):
    words = set()
    for tl in similarTaglines:
        for word in tl.split(' '):
            if word: words.add(word)
    return list(words)

def taglineToBody(taglines : list[str]):
    bodies = []
    for tl in taglines:
        body = df.row(by_predicate=(pl.col("tagline") == tl))[-1]
        bodies.append(body)
    return bodies

def taglineToCard(taglines : list[str]):
    cards = []
    for tl in taglines:
        card = df.row(by_predicate=(pl.col("tagline") == tl))[0]
        cards.append(card)
    return cards

#####################################################################################################################

##### Parse Article #####
    
url = "https://thehill.com/policy/healthcare/public-global-health/485602-virus-expert-as-much-as-70-percent-of-worlds/"

driver.get(url)
WebDriverWait(driver, 10)
paragraphs = [p.text for p in driver.find_elements(By.TAG_NAME, 'p') if not p.text.isspace() and p.text and 'Copyright' not in p.text]
print(paragraphs)
articleTitle = driver.find_element(By.TAG_NAME, 'h1').text
driver.quit()

# Get dataset
df = data(numGroups=20)

# Our tagline
tagline = "Coronavirus is inevitable but not existential"
print(f'Our tagline: {tagline}')
tagline_words = parseTagline(tagline)
parsed_tl = ' '.join(tagline_words)

############### USING DATASET ##################

tl_df = df.get_columns()[1].to_list()
parsed_tl_df = [' '.join(parseTagline(tl)).lower() for tl in tl_df]

"""Search through dataframe for similar taglines""" 
# all_synonym_tls = findSimilarTaglinesTFIDF(parsed_tl, parsed_tl_df)
all_synonym_tls = findSimilarTaglinesJaccard(parsed_tl, tl_df)
PERCENTILE = 99
perc_tls = abovePercentile(all_synonym_tls, PERCENTILE)

"""Get corresponding bodies of each similar tagline"""
bodies = taglineToBody(perc_tls)
cards = taglineToCard(perc_tls)
print(cards)

""" != None or run.bold or 'Bold' in run.style.name or run.style.name == 'Emphasis' or run.style.name == 'Style Bold Underline'"""

impt_words : dict[ str : float ] = {}
for body in bodies:
    # print('Body text: ', body.text)
    runs = body.runs
    for run in runs:
        # print(run.text)
        if run.text and run.font.highlight_color != None: 
            for word in parseTagline(run.text):
                if word in impt_words:
                    impt_words[word] += 1
                else:
                    impt_words[word] = 1
for tl in perc_tls:
    for word in parseTagline(tl):
        if word in impt_words:
            impt_words[word] += 1
        else:
            impt_words[word] = 2
for word in tagline_words:
    if word in impt_words:
        impt_words[word] += 1
    else:
        impt_words[word] = 3
print(impt_words)

weighted_words = {word : np.log1p(val) for word, val in impt_words.items()}
print(weighted_words)

# similar_taglines = list(perc_tls.keys())
# unique_tl_words = uniqueTaglineWords(similar_taglines)
# print(unique_tl_words)


normalEvidence = unlabeledRanking(list(impt_words.keys()), paragraphs)
weightedEvidence = weightedRanking(impt_words, paragraphs)

# finalEvidence = {}
# for sentence, val in rawEvidence.items():
#     if sentence in finalEvidence:
#         finalEvidence[sentence] += val
#     else:
#         finalEvidence[sentence] = val
# print(finalEvidence)

# enhancedRankings = sortSentences(finalSentences, paragraphs)
# print('Enhanced rankings: ', list(finalSentences.values()))
# print('Enhanced rankings: ', np.round(np.log(list(finalSentences.values()))))

makeDoc(tagline, 'PressTV', articleTitle, 'May 25, 2015', normalEvidence, 'raw')
makeDoc(tagline, 'PressTV', articleTitle, 'May 25, 2015', weightedEvidence, 'weighted')

generateTxt(weighted_words, 'words')
generateTxt(weightedEvidence, 'Wevidence')
generateTxt(normalEvidence, 'Revidence')

# ###################################################

# ############ ORIGINAL TAGLINE #################
# # normalRankings = unlabeledRanking(tagline_words, paragraphs)
# # print('Normal rankings: ', list(normalRankings.values()))
# # print('Normal rankings: ', np.round(np.log(list(normalRankings.values()))))
# # makeDoc(tagline, 'PressTV', articleTitle, 'May 25, 2015', normalRankings, 'card1')


